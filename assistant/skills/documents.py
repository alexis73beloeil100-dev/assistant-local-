"""Ecrire des documents : texte, Word, PDF.

content.py savait LIRE le texte, les PDF, les Word, les Excel et les
PowerPoint. Il ne savait rien ecrire. "Fais-moi un compte rendu" se terminait
donc par un texte affiche a l'ecran, que l'utilisateur recopiait a la main
dans Word -- ce qui annule l'interet d'avoir demande.

Les bibliotheques etaient deja la, embarquees pour la lecture : python-docx
et reportlab. Il ne manquait que le chemin inverse.

Deux regles, pour les memes raisons qu'ailleurs dans ce projet :

  1. On n'ecrase JAMAIS un fichier existant sans le dire. Un document qu'on
     croyait fini et qu'une phrase mal comprise remplace ne se recupere pas,
     et l'utilisateur ne s'en apercoit qu'en le rouvrant.
  2. Les chemins proteges refusent, meme sur confirmation.
"""
from __future__ import annotations

from pathlib import Path

from assistant import safety

FORMATS = ("txt", "md", "docx", "pdf")

# Marge et corps de texte du PDF, en points typographiques (72 par pouce).
MARGE = 56
INTERLIGNE = 15
TAILLE = 11


def _ecrire_texte(chemin: Path, titre: str, texte: str) -> None:
    contenu = f"{titre}\n{'=' * len(titre)}\n\n{texte}\n" if titre else texte
    chemin.write_text(contenu, encoding="utf-8")


def _ecrire_docx(chemin: Path, titre: str, texte: str) -> None:
    from docx import Document

    document = Document()
    if titre:
        document.add_heading(titre, level=1)
    # Une ligne vide separe deux paragraphes : c'est la convention de tout
    # texte dicte. Sans ce decoupage, Word recevrait un seul bloc illisible.
    for bloc in texte.split("\n\n"):
        if bloc.strip():
            document.add_paragraph(bloc.strip())
    document.save(str(chemin))


def _ecrire_pdf(chemin: Path, titre: str, texte: str) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(
        str(chemin), pagesize=A4,
        leftMargin=MARGE, rightMargin=MARGE,
        topMargin=MARGE, bottomMargin=MARGE)

    elements = []
    if titre:
        elements.append(Paragraph(titre, styles["Title"]))
        elements.append(Spacer(1, INTERLIGNE))
    for bloc in texte.split("\n\n"):
        if bloc.strip():
            # Les chevrons sont du balisage pour reportlab : un texte qui en
            # contient ferait echouer la generation au lieu de s'afficher.
            propre = (bloc.strip().replace("&", "&amp;")
                      .replace("<", "&lt;").replace(">", "&gt;"))
            elements.append(Paragraph(propre.replace("\n", "<br/>"),
                                      styles["BodyText"]))
            elements.append(Spacer(1, 6))
    document.build(elements)


ECRIVAINS = {
    "txt": _ecrire_texte,
    "md": _ecrire_texte,
    "docx": _ecrire_docx,
    "pdf": _ecrire_pdf,
}


def ecrire(chemin: str, texte: str, titre: str = "", ask=None) -> str:
    """Ecrit un document. Le format est deduit de l'extension.

    Le format vient de l'extension et n'est jamais devine autrement : on
    ecrirait un PDF dans un fichier .docx que Word refuserait d'ouvrir, sans
    que rien n'ait signale d'erreur au moment de l'ecriture.
    """
    if not str(texte).strip():
        return "Rien a ecrire : le texte est vide."

    cible = Path(chemin).expanduser()
    extension = cible.suffix.lstrip(".").lower()
    if extension not in ECRIVAINS:
        return (f"Je ne sais pas ecrire de \".{extension}\". "
                f"Formats disponibles : {', '.join(FORMATS)}.")

    if safety.is_protected(str(cible)):
        return (f"{cible} est dans les chemins proteges. Je n'y ecris pas, "
                "meme sur confirmation.")

    existant = cible.exists()
    action = safety.Action(
        kind="fichier",
        summary=("REMPLACER " if existant else "Ecrire ") + str(cible),
        targets=[str(cible)],
        # Creer un fichier se defait en le supprimant. En remplacer un se
        # defait rarement : l'ancien contenu n'existe plus nulle part.
        reversible=not existant,
        details=(f"{len(texte)} caracteres, format {extension}"
                 + ("\n    Le fichier existe deja et son contenu actuel sera "
                    "perdu." if existant else "")),
        routine=not existant,
    )
    try:
        safety.guard(action, ask=ask)
    except safety.Refused as exc:
        return str(exc)

    try:
        cible.parent.mkdir(parents=True, exist_ok=True)
        ECRIVAINS[extension](cible, titre.strip(), str(texte))
    except (OSError, ImportError, ValueError) as exc:
        return f"Ecriture impossible : {type(exc).__name__}: {exc}"

    poids = cible.stat().st_size
    verbe = "remplace" if existant else "cree"
    return f"{cible} {verbe} ({poids} octets)."
