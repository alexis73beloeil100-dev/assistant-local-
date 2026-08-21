"""Lecture du contenu des fichiers, a la demande et sans rien conserver.

Principe, qui prolonge celui de l'index : on lit au moment ou la question est
posee, le texte vit en memoire le temps de la reponse, puis il disparait.
Aucun cache, aucun extrait recopie sur le disque.

L'index des noms sert a choisir QUELS fichiers ouvrir. Sans lui il faudrait
parcourir 914 000 fichiers a chaque question ; avec lui on en ouvre quelques
dizaines, ciblees.
"""
from __future__ import annotations

import concurrent.futures
import os
import zipfile
from pathlib import Path

from assistant.index import db
from assistant.util import human_size, norm

# --- Ce qu'on sait lire -----------------------------------------------------

TEXT_EXT = {
    "txt", "md", "log", "csv", "tsv", "json", "xml", "yaml", "yml", "toml",
    "ini", "cfg", "conf", "properties", "env", "reg", "srt", "vtt",
    "py", "js", "ts", "jsx", "tsx", "c", "h", "cpp", "hpp", "cs", "java",
    "go", "rs", "rb", "php", "lua", "sh", "bat", "ps1", "sql", "r", "m",
    "html", "htm", "css", "scss", "vue", "svelte",
    "gitignore", "gitattributes", "editorconfig", "dockerfile", "makefile",
    "spec", "in", "cmake", "gradle", "pro", "pri", "def", "map", "list",
    "nfo", "diz", "asc", "pem", "crt", "url", "desktop", "lnk_txt",
    # Fichiers de configuration de jeux : c'est souvent ce qu'on cherche.
    "acd", "kn5", "ini", "sav", "cfg", "settings",
}

RICH_EXT = {"pdf", "docx", "xlsx", "pptx"}

# Au-dela, on lit le debut seulement : personne ne veut 400 Mo de log dans
# une fenetre de contexte, et le modele n'en ferait rien de bon.
MAX_READ_BYTES = 2 * 1024 * 1024
MAX_CHARS_DEFAULT = 20_000

# Garde-fous de la recherche dans les contenus. 3000 fichiers se lisent en
# quelques secondes sur 8 threads : le plafond precedent (400) etait trop bas
# et laissait passer a cote du fichier cherche.
MAX_FILES_SCANNED = 3000
MAX_FILE_SIZE_SCAN = 8 * 1024 * 1024
WORKERS = 8


def kind(path: str) -> str:
    """text, rich, ou binaire."""
    name = Path(path).name.lower()
    ext = name.rpartition(".")[2] if "." in name else name
    if ext in RICH_EXT:
        return "rich"
    if ext in TEXT_EXT or name in TEXT_EXT:
        return "text"
    return "binaire"


# --- Extraction -------------------------------------------------------------

def _read_text(path: str, max_bytes: int) -> str:
    """Lit un fichier texte en devinant son encodage.

    Windows melange UTF-8, UTF-8 BOM et cp1252 dans les memes dossiers. Un
    decodage strict echouerait sur un fichier sur trois.
    """
    with open(path, "rb") as fh:
        raw = fh.read(max_bytes)

    for encoding in ("utf-8-sig", "utf-8"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            pass

    try:
        from charset_normalizer import from_bytes

        best = from_bytes(raw).best()
        if best:
            return str(best)
    except Exception:  # noqa: BLE001
        pass

    return raw.decode("cp1252", errors="replace")


def _read_pdf(path: str, max_pages: int = 40) -> str:
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages = []
    for i, page in enumerate(reader.pages):
        if i >= max_pages:
            pages.append(f"[... {len(reader.pages) - max_pages} pages non lues]")
            break
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _read_docx(path: str) -> str:
    import docx

    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_xlsx(path: str, max_rows: int = 300) -> str:
    import openpyxl

    book = openpyxl.load_workbook(path, read_only=True, data_only=True)
    out = []
    for sheet in book.worksheets:
        out.append(f"--- feuille : {sheet.title} ---")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i >= max_rows:
                out.append(f"[... suite tronquee]")
                break
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    book.close()
    return "\n".join(out)


def _read_pptx(path: str) -> str:
    """Extrait le texte d'un .pptx sans dependance supplementaire.

    Un pptx est un zip de XML : on y pioche les noeuds de texte plutot que
    d'ajouter python-pptx pour ce seul usage.
    """
    import re

    out = []
    with zipfile.ZipFile(path) as z:
        slides = sorted(n for n in z.namelist()
                        if n.startswith("ppt/slides/slide") and n.endswith(".xml"))
        for name in slides:
            xml = z.read(name).decode("utf-8", errors="replace")
            texts = re.findall(r"<a:t>(.*?)</a:t>", xml, re.DOTALL)
            if texts:
                out.append(f"--- {Path(name).stem} ---")
                out.extend(texts)
    return "\n".join(out)


def extract(path: str, max_chars: int = MAX_CHARS_DEFAULT) -> tuple[bool, str]:
    """Rend (succes, texte) pour n'importe quel fichier lisible."""
    if not os.path.isfile(path):
        return False, f"{path} n'existe pas."

    try:
        size = os.path.getsize(path)
    except OSError as exc:
        return False, f"Illisible : {exc}"

    ext = Path(path).suffix.lower().lstrip(".")
    try:
        if ext == "pdf":
            text = _read_pdf(path)
        elif ext == "docx":
            text = _read_docx(path)
        elif ext == "xlsx":
            text = _read_xlsx(path)
        elif ext == "pptx":
            text = _read_pptx(path)
        else:
            text = _read_text(path, MAX_READ_BYTES)
    except Exception as exc:  # noqa: BLE001
        return False, f"Lecture impossible ({type(exc).__name__}: {exc})."

    if not text.strip():
        return False, (
            f"{Path(path).name} ne contient pas de texte lisible "
            f"({human_size(size)}, probablement binaire)."
        )

    # Un fichier binaire lu comme du texte donne une bouillie de caracteres de
    # controle : on le detecte plutot que de la servir au modele.
    control = sum(1 for c in text[:4000] if ord(c) < 9 or 13 < ord(c) < 32)
    if control > len(text[:4000]) * 0.08:
        return False, (
            f"{Path(path).name} est un fichier binaire, son contenu n'est pas "
            "du texte."
        )

    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars]
    return True, text + ("\n\n[... suite tronquee]" if truncated else "")


# --- Outils exposes ---------------------------------------------------------

def read(path: str, max_chars: int = MAX_CHARS_DEFAULT) -> str:
    """Lit un fichier et rend son contenu. Rien n'est conserve."""
    # Une image n'a pas de texte a extraire au sens habituel : elle passe par
    # la reconnaissance visuelle, pas par le decodage de caracteres.
    from assistant.skills import vision

    if vision.is_image(path):
        return vision.read_image(path)

    ok, text = extract(path, max_chars)
    if not ok:
        return text
    size = human_size(os.path.getsize(path))
    return f"--- {path} ({size}) ---\n{text}"


def _candidates(scope: str | None, ext: str | None, name_hint: str | None,
                limit: int) -> list[str]:
    """Choisit les fichiers a ouvrir, via l'index des noms."""
    where = ["is_dir = 0", "is_cache = 0", "size > 0", "size <= ?"]
    params: list = [MAX_FILE_SIZE_SCAN]

    if scope:
        base = scope.rstrip("\\") + "\\"
        where.append("path >= ? AND path < ?")
        params.extend([base, base + "￿"])
    if ext:
        where.append("ext = ?")
        params.append(ext.lower().lstrip("."))
    else:
        readable = tuple(TEXT_EXT | RICH_EXT)
        where.append(f"ext IN ({','.join('?' * len(readable))})")
        params.extend(readable)
    if name_hint:
        where.append("name LIKE ?")
        params.append(f"%{name_hint}%")

    sql = (f"SELECT path FROM files WHERE {' AND '.join(where)} "
           f"ORDER BY mtime DESC LIMIT ?")
    params.append(limit)

    with db.connect() as conn:
        return [r["path"] for r in conn.execute(sql, params).fetchall()]


def _scan_one(path: str, needle: str) -> tuple[str, list[str]] | None:
    ok, text = extract(path, max_chars=MAX_READ_BYTES)
    if not ok:
        return None
    lines = text.splitlines()
    hits = [f"    l.{i + 1}: {line.strip()[:160]}"
            for i, line in enumerate(lines) if needle in line.lower()]
    return (path, hits[:5]) if hits else None


def search_in_files(texte: str, dossier: str | None = None,
                    ext: str | None = None, nom: str | None = None,
                    limit: int = 12) -> str:
    """Cherche une expression DANS le contenu des fichiers.

    Restreins toujours avec un dossier, une extension ou un morceau de nom :
    ouvrir des milliers de fichiers a chaque question serait interminable.
    """
    if not db.is_ready():
        return ("L'index des noms n'est pas encore pret, il sert a choisir "
                "quels fichiers ouvrir. Reessaie dans un instant.")
    if not texte.strip():
        return "Precise ce qu'il faut chercher."

    paths = _candidates(dossier, ext, nom, MAX_FILES_SCANNED)
    if not paths:
        return ("Aucun fichier lisible ne correspond a ce filtre. "
                "Elargis le dossier ou l'extension.")

    needle = texte.lower()
    found = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=WORKERS) as pool:
        for result in pool.map(lambda p: _scan_one(p, needle), paths):
            if result:
                found.append(result)
            if len(found) >= limit:
                break

    if not found:
        return (f"\"{texte}\" n'apparait dans aucun des {len(paths)} fichiers "
                "examines.")

    out = [f"\"{texte}\" trouve dans {len(found)} fichier(s) "
           f"sur {len(paths)} examines :"]
    for path, hits in found:
        out.append(f"  {path}")
        out.extend(hits)
    return "\n".join(out)


def peek(dossier: str, limit: int = 10) -> str:
    """Apercu d'un dossier : ce qu'il contient, sans tout ouvrir."""
    if not os.path.isdir(dossier):
        return f"{dossier} n'est pas un dossier."
    try:
        entries = sorted(os.scandir(dossier), key=lambda e: e.name.lower())
    except OSError as exc:
        return f"Illisible : {exc}"

    dirs = [e.name for e in entries if e.is_dir()][:limit]
    files = [e for e in entries if e.is_file()][:limit]

    out = [f"--- {dossier} ---"]
    if dirs:
        out.append("Dossiers : " + ", ".join(dirs))
    for e in files:
        try:
            taille = human_size(e.stat().st_size)
        except OSError:
            taille = "?"
        out.append(f"  {taille:>10}  {e.name}  [{kind(e.path)}]")
    return "\n".join(out)
